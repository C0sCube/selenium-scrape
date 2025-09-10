import base64
from io import BytesIO, StringIO
from typing import List, Dict, Any
import fitz, pdfplumber
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document


class Parser:
    @staticmethod
    def __parse_entry(entry: Dict[str, Any]) -> Dict[str, Any]:
        """
        Parse a single response entry (pdf/html/table_html/text).
        Returns a dict with:
          - type: content_type
          - content: ordered list of {"type":"text","text":...} or {"type":"table","table":DataFrame}
          - tables: list of detected DataFrames (for convenience)
          - raw_text: concatenated text-only content
        """
        content_type = entry.get("type", "str")
        raw_content = entry.get("value", "")
        result = {"type": content_type, "content": [], "tables": [], "raw_text": ""}

        def clean_cell(cell):
            if cell is None:
                return ""
            cleaned = "".join(str(cell).splitlines())
            cleaned = " ".join(cleaned.split())
            return cleaned.strip()

        try:
            if content_type == "pdf":
                # decode base64 and prepare streams
                pdf_bytes = base64.b64decode(raw_content)
                pdf_file = BytesIO(pdf_bytes)

                with pdfplumber.open(pdf_file) as plumber_pdf:
                    fitz_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")

                    # iterate pages
                    for page_num, plumber_page in enumerate(plumber_pdf.pages):
                        fitz_page = fitz_pdf[page_num]

                        elements = []  # will collect dicts: {type: 'text'/'table', top: float, left: float, ...}

                        # --- 1) detect tables via pdfplumber.find_tables() (gives bboxes + extract)
                        table_objs = plumber_page.find_tables()
                        table_bboxes = []
                        for t in table_objs:
                            bbox = t.bbox  # (x0, top, x1, bottom)
                            table_bboxes.append(bbox)
                            try:
                                raw_table = t.extract()  # list-of-rows
                            except Exception:
                                raw_table = plumber_page.extract_table(t.settings) if hasattr(plumber_page, 'extract_table') else []
                            # convert to DataFrame, clean cells
                            if raw_table:
                                cleaned_table = []
                                for row in raw_table:
                                    cleaned_row = [clean_cell(cell) for cell in row]
                                    cleaned_table.append(cleaned_row)
                                df = pd.DataFrame(cleaned_table)
                                elements.append({
                                    "type": "table",
                                    "table": df,
                                    "top": float(bbox[1]),
                                    "left": float(bbox[0]),
                                    "page": page_num
                                })
                                result["tables"].append(df)

                        # --- 2) extract words via fitz and group into lines
                        # fitz words typically: (x0, y0, x1, y1, "word", block_no, line_no, word_no)
                        words = fitz_page.get_text("words")  # may be empty list
                        if words:
                            # sort by y then x (top to bottom, left to right)
                            words_sorted = sorted(words, key=lambda w: (round(w[1], 2), round(w[0], 2)))
                            # group words into lines by y coordinate tolerance
                            lines = []
                            current_line = []
                            current_y = None
                            y_tol = 3.0  # tuning parameter (points). Increase if lines split too often.
                            for w in words_sorted:
                                # defensive unpack: ensure at least 5 elements
                                if len(w) >= 5:
                                    x0, y0, x1, y1, text = float(w[0]), float(w[1]), float(w[2]), float(w[3]), str(w[4])
                                else:
                                    continue

                                if current_y is None:
                                    current_y = y0
                                    current_line = [(x0, text, y0, y1, x1)]
                                elif abs(y0 - current_y) <= y_tol:
                                    current_line.append((x0, text, y0, y1, x1))
                                else:
                                    # flush current line
                                    current_line.sort(key=lambda c: c[0])
                                    line_text = " ".join([c[1] for c in current_line]).strip()
                                    top = min(c[2] for c in current_line)
                                    left = min(c[0] for c in current_line)
                                    right = max(c[4] for c in current_line)
                                    bottom = max(c[3] for c in current_line)
                                    lines.append({
                                        "text": line_text,
                                        "top": float(top),
                                        "left": float(left),
                                        "right": float(right),
                                        "bottom": float(bottom),
                                        "page": page_num
                                    })
                                    # start new line
                                    current_y = y0
                                    current_line = [(x0, text, y0, y1, x1)]
                            # flush last line
                            if current_line:
                                current_line.sort(key=lambda c: c[0])
                                line_text = " ".join([c[1] for c in current_line]).strip()
                                top = min(c[2] for c in current_line)
                                left = min(c[0] for c in current_line)
                                right = max(c[4] for c in current_line)
                                bottom = max(c[3] for c in current_line)
                                lines.append({
                                    "text": line_text,
                                    "top": float(top),
                                    "left": float(left),
                                    "right": float(right),
                                    "bottom": float(bottom),
                                    "page": page_num
                                })

                            # --- 3) classify each line whether it intersects a detected table bbox
                            def intersects(b1, b2):
                                # b1 = (left, top, right, bottom)
                                # b2 similar
                                return not (b1[2] < b2[0] or b1[0] > b2[2] or b1[3] < b2[1] or b1[1] > b2[3])

                            for ln in lines:
                                line_bbox = (ln["left"], ln["top"], ln["right"], ln["bottom"])
                                in_table = False
                                for tb in table_bboxes:
                                    if intersects(line_bbox, tb):
                                        in_table = True
                                        break
                                elements.append({
                                    "type": "text",
                                    "text": ln["text"],
                                    "top": ln["top"],
                                    "left": ln["left"],
                                    "page": page_num,
                                    "label": "table_text" if in_table else "body_text"
                                })

                        # --- 4) sort all elements on this page by (top, left) so visual order preserved
                        elements_on_page = [el for el in elements if el.get("page") == page_num]
                        elements_on_page.sort(key=lambda x: (float(x.get("top", 0.0)), float(x.get("left", 0.0))))

                        # add to result.content preserving order
                        for el in elements_on_page:
                            if el["type"] == "text":
                                # skip empty lines
                                if el["text"].strip():
                                    result["content"].append({"type": "text", "text": el["text"]})
                            elif el["type"] == "table":
                                result["content"].append({"type": "table", "table": el["table"]})

                    # finalize raw_text
                    result["raw_text"] = "\n".join([el["text"] for el in result["content"] if el["type"] == "text"])

            elif content_type == "html":
                soup = BeautifulSoup(raw_content, "html.parser")
                text = soup.get_text(separator="\n").strip()
                result["content"].append({"type": "text", "text": text})
                result["raw_text"] = text

            elif content_type == "table_html":
                df = pd.read_html(StringIO(raw_content))[0]
                result["content"].append({"type": "table", "table": df})
                result["tables"].append(df)

            else:
                result["content"].append({"type": "text", "text": f"[Unsupported Datatype: {content_type}]"})

        except Exception as e:
            result["content"].append({"type": "text", "text": f"[Failed to parse: {e}]"})
        return result

    @classmethod
    def generate_cache_doc_report(cls, comparison_json: Dict[str, Any], output_path="DepositRate_Comparison_Report.docx") -> str:
        """
        Builds a Word document that preserves the sequence of parsed content.
        """
        document = Document()
        sorted_records = comparison_json.get("records", [])
        metadata = comparison_json.get("metadata", {})

        document.add_heading(">>>Metadata Overview<<<", level=1)
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Key'
        hdr_cells[1].text = 'Value'

        for key, value in metadata.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

        document.add_paragraph("\n")

        for record in sorted_records:
            bank_name = record.get("bank_name", "")
            bank_code = record.get("bank_code", "")
            scraped_data = record.get("scraped_data", [])

            document.add_heading(f">> {bank_code} : {bank_name}", level=2)
            document.add_paragraph("================================================")

            for scrape in scraped_data:
                responses = scrape.get("response", [])
                action = scrape.get("action", "")
                timestamp = scrape.get("timestamp", "")
                webpage = scrape.get("webpage", "")
                data_present = scrape.get("data_present", "")
                response_count = scrape.get("response_count", "")

                document.add_heading(
                    f"Action: {action} | Timestamp: {timestamp} | Present: {str(data_present)} | Count: {response_count if data_present else 0}",
                    level=3
                )
                document.add_paragraph(f"Website: {webpage}")

                for response_entry in responses:
                    titles = response_entry.get("title", [])
                    titles = [titles] if isinstance(titles, str) else titles

                    parsed = cls.__parse_entry(response_entry)
                    content_stream = parsed["content"]

                    document.add_paragraph("---------------------------")
                    for idx, line in enumerate(titles):
                        document.add_paragraph(f"TITLE {idx+1}: {line}")

                    for item in content_stream:
                        if item["type"] == "text":
                            document.add_paragraph(item["text"])
                        elif item["type"] == "table":
                            df = item["table"]
                            if df.empty:
                                continue
                            # create docx table
                            cols = len(df.columns)
                            table = document.add_table(rows=1, cols=cols)
                            table.style = 'Table Grid'
                            for i, col_name in enumerate(df.columns):
                                table.cell(0, i).text = str(col_name)
                            for _, row in df.iterrows():
                                row_cells = table.add_row().cells
                                for i, val in enumerate(row):
                                    row_cells[i].text = str(val)
                            document.add_paragraph("")  # spacing

                    document.add_paragraph("")

            document.add_page_break()

        document.save(output_path)
        return output_path

    @classmethod
    def generate_cache_excel_report(cls, comparison_json: Dict[str, Any], output_path="DepositRate_Comparison_Report.xlsx") -> str:
        """
        Concatenate headers and parsed content into one sheet per bank.
        """
        sorted_records = comparison_json.get("records", [])
        metadata = comparison_json.get("metadata", {})

        with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
            meta_df = pd.DataFrame(list(metadata.items()), columns=["Key", "Value"])
            meta_df.to_excel(writer, sheet_name="Metadata", index=False)

            for record in sorted_records:
                bank_name = record.get("bank_name", "UnknownBank")
                bank_code = record.get("bank_code", "")
                scraped_data = record.get("scraped_data", [])
                sheet_name = f"{bank_code}_{bank_name}"[:31]

                bank_frames: List[pd.DataFrame] = []

                for scrape in scraped_data:
                    responses = scrape.get("response", [])
                    action = scrape.get("action", "")
                    timestamp = scrape.get("timestamp", "")
                    webpage = scrape.get("webpage", "")
                    data_present = scrape.get("data_present", "")
                    response_count = scrape.get("response_count", "")

                    for response_entry in responses:
                        titles = response_entry.get("title", [])
                        titles = [titles] if isinstance(titles, str) else titles

                        parsed = cls.__parse_entry(response_entry)
                        content_stream = parsed["content"]

                        header_rows = [
                            [f"Action: {action}", f"Timestamp: {timestamp}", f"Present: {data_present}", f"Count: {response_count}"],
                            [f"Website: {webpage}"],
                            ["---------------------------"],
                        ]
                        for idx, title in enumerate(titles):
                            header_rows.append([f"TITLE {idx+1}: {title}"])

                        bank_frames.append(pd.DataFrame(header_rows))

                        for item in content_stream:
                            if item["type"] == "text":
                                bank_frames.append(pd.DataFrame([[item["text"]]]))
                            elif item["type"] == "table":
                                bank_frames.append(item["table"])

                        bank_frames.append(pd.DataFrame([[""]]))  # Spacer

                if bank_frames:
                    final_df = pd.concat(bank_frames, ignore_index=True)
                else:
                    final_df = pd.DataFrame([["No data captured"]])

                # write
                final_df.to_excel(writer, sheet_name=sheet_name, index=False)

        return output_path
