import time ,re,os, hmac, hashlib, inspect, dateutil, base64, pdfplumber, ocrmypdf, tempfile,shutil, tabula
import pandas as pd
from bs4 import BeautifulSoup
from dateutil.parser import parse
from io import StringIO, BytesIO
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import PatternFill
from docx import Document
from docx.shared import Pt
from app.utils import Helper


class OperationExecutor:
    
    
    cache_doc_name = ""
    
    def __init__(self, logger = None):
        
        self.logger = logger
        self.procedures = {
            "ext_date": self.extract_date,
            "sha256": self._generate_hash_sha256,
            "sha1": self._generate_hash_sha1,
            "normalize_df": self._generalize_table_df,
            "original":self._boomerang
        }
        
        self.type_compatibility = {
            "normalize_df": ["table_html"],
            "sha1": ["table_html", "html", "pdf"],
            "sha256": ["table_html", "html", "pdf"],
            "ext_date": ["html"],
            "original": ["pdf", "html", "table_html"]
        }

    #============ PROCEDURES ==============
    @staticmethod
    def extract_date(text: str) -> str:
        if not isinstance(text, str):
            print("Invalid data type. Expected string.")
            return ""
        
        output_format = "%Y%m%d"
        date_patterns = [r"(\d{2}[.\-/]+\d{2}[.\-/]+\d{4}",
                        r"\d{1,2}\s*(?:th|st|rd|nd)\s*[A-Za-z]+\s*\d{4}",
                        r"\d{2}[\.\-\/]+[A-Za-z]+[\.\-\/]+\d{4}",
                        r"\d{2}\s*[A-Za-z]+\s*\d{4})"]
        matches = re.findall(r"|".join(date_patterns), text, re.IGNORECASE)
        
        if matches:
            date_str = " ".join(matches)
            try:
                dt_object = parse(date_str, fuzzy=True)
                return dt_object.strftime(output_format)
            except dateutil.parser._parser.ParserError as e:
                print(f"[ERROR]: {e}")
                return date_str
        return text
    
    def _boomerang(self,data):
        return data
    
    def _generate_hash_sha256(self,text:str)->str:
        if not isinstance(text,str):
            return f"{inspect.currentframe().f_code.co_name}: input non str"
        return hashlib.sha256(text.encode()).hexdigest()

    def _generate_hash_sha1(self,text:str)->str:
        if not isinstance(text,str):
            return f"{inspect.currentframe().f_code.co_name}: input non str"
        return hashlib.sha1(text.encode()).hexdigest()
    
    def _generalize_table_df(self,html_str)->str:
        MAX_COLUMN=15
        cols = [f"column_{i}" for i in range(1, MAX_COLUMN + 1)]
        dfs = pd.read_html(StringIO(html_str), flavor='html5lib')
        if not dfs:
            return pd.DataFrame(columns=cols)
        
        all_dfs = []
        for df in dfs:
            if df.empty:
                continue

            df = df.iloc[:, :MAX_COLUMN].copy()  # truncate if too many columns
            df.columns = cols[:df.shape[1]]      # rename existing columns
            for i in range(df.shape[1], MAX_COLUMN):
                df[cols[i]] = ""                 # fill missing columns with empty strings

            df = df.reindex(columns=cols)        # ensure consistent column order
            all_dfs.append(df)

        final_df = pd.concat(all_dfs, ignore_index=True) if all_dfs else pd.DataFrame(columns=cols)
        
        norm_df = final_df.to_csv(index=False,header=False, sep='|', lineterminator='\n')
        return norm_df


    
    @staticmethod
    def save_tables_to_excel(tables, output_dir="tables", output_file="all_tables.xlsx", consolidate_save=True):
        """
        Save tables to Excel.
        If separator=True, saves all tables in one file with multiple sheets.
        If separator=False, saves each table in a separate Excel file.
        """
        if consolidate_save:
            full_path = os.path.join(output_dir, output_file)
            with pd.ExcelWriter(full_path, engine="openpyxl") as writer:
                for i, table in enumerate(tables):
                    df = pd.read_html(str(table))[0]
                    sheet_name = f"Table_{i+1}"
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
                    # print(f"Added sheet: {sheet_name}")
            # return full_path
        else:
            saved_files = []
            for i, table in enumerate(tables):
                df = pd.read_html(str(table))[0]
                file_path = os.path.join(output_dir, f"table_{i+1}.xlsx")
                df.to_excel(file_path, index=False)
                saved_files.append(file_path)
                # print(f"Saved: {file_path}")
            # return saved_files

    @staticmethod
    def save_tables_html(tables,output_dir="output_html",output_file="combined.html",separator=None,wrap_html=True):
        
        if isinstance(tables,str):tables = [tables]
        if separator is None:
            # Save each table separately
            for i, table in enumerate(tables):
                filename = f"content_{i+1}.html"
                path = os.path.join(output_dir, filename)
                with open(path, "w", encoding="utf-8") as f:
                    if wrap_html:
                        f.write("<html><body>\n")
                    f.write(str(table) + "\n")
                    if wrap_html:
                        f.write("</body></html>")
        else:
            # Save all tables in one file with separator
            path = os.path.join(output_dir, output_file)
            with open(path, "w", encoding="utf-8") as f:
                if wrap_html:
                    f.write("<html><body>\n")
                for i, table in enumerate(tables):
                    f.write(str(table) + "\n")
                    if i < len(tables) - 1:
                        f.write(separator + "\n")
                if wrap_html:
                    f.write("</body></html>")
  
    #Core Functionality
    def runner(self, data, function_to_execute):
        p_dict = data.copy()
        records = p_dict.get("records", [])

        for record in records:
            # self.logger.info(f"Processing : {record['bank_name']}")
            print(f">>Processing {record['bank_name']}")
            response_data = record.get("scraped_data", [])
            if not response_data:
                continue

            new_scraped_data = []

            for action in response_data:
                if not action.get("data_present"):
                    continue

                response = action.get("response")
                if not response:
                    continue

                for _packet_ in response:
                    check_packet = _packet_.copy()

                    for stage_name, operations in function_to_execute.items():
                        for operation in operations:
                            # Unpack operation with optional expected_type
                            if len(operation) == 4:
                                func_name, source_key, target_key, expected_type = operation
                            else:
                                func_name, source_key, target_key = operation
                                expected_type = None

                            if target_key in check_packet:
                                raise ValueError(
                                    f"`target_key` cannot be similar to any of these keys: {list(check_packet.keys())}"
                                )

                            func = self.procedures.get(func_name)
                            if not func:
                                raise ValueError(f"Function '{func_name}' not found in procedures.")

                            input_value = _packet_.get(source_key)
                            if input_value is None:
                                continue

                            # Apply type check only in primary stage
                            if stage_name == "primary" and expected_type:
                                packet_type = _packet_.get("type", "").lower()
                                if isinstance(expected_type, list):
                                    if packet_type not in [t.lower() for t in expected_type]:
                                        continue
                                elif packet_type != expected_type.lower():
                                    continue

                            _packet_[target_key] = func(input_value)

                    new_scraped_data.append(_packet_)

            record["scraped_data"] = new_scraped_data

        return p_dict
        
    def process_comparison(self, old_json: dict, new_json: dict, key: str = "hash256") -> dict:
    
        def __extract_scraped_items(data, bank_code):
            for record in data.get("records", []):
                if record.get("bank_code") == bank_code:
                    return [entry for entry in record.get("scraped_data", []) if key in entry]
            return []

        def __build_comparison_result(result):
            return {
                "comparison_result": {
                    "new": result["new_packets"],
                    "removed": result["removed_packets"],
                    "unchanged": list(result["unchanged_keys"]),
                    "summary": {
                        "old_total": result["old_total"],
                        "new_total": result["new_total"],
                        "new_count": len(result["new_packets"]),
                        "removed_count": len(result["removed_packets"]),
                        "unchanged_count": len(result["unchanged_keys"])
                    }
                }
            }

        for new_record in new_json.get("records", []):
            bank_code = new_record.get("bank_code")
            old_items = __extract_scraped_items(old_json, bank_code)
            new_items = __extract_scraped_items(new_json, bank_code)

            old_keys = {item[key] for item in old_items}
            new_keys = {item[key] for item in new_items}

            new_only_keys = new_keys - old_keys
            removed_keys = old_keys - new_keys
            unchanged_keys = old_keys & new_keys

            result = {
                "key": key,
                "new_keys": new_only_keys,
                "removed_keys": removed_keys,
                "unchanged_keys": unchanged_keys,
                "new_packets": [item for item in new_items if item[key] in new_only_keys],
                "removed_packets": [item for item in old_items if item[key] in removed_keys],
                "old_total": len(old_items),
                "new_total": len(new_items)
            }

            new_record.update(__build_comparison_result(result))
            new_record.pop("scraped_data", None)

        return new_json

    #Report Generation
    
    def _parse_table(self, entry):
        import base64
        import tempfile
        from io import BytesIO, StringIO
        import pdfplumber
        import pandas as pd
        from bs4 import BeautifulSoup
        import ocrmypdf

        content_type = entry.get("type", "str")
        raw_content = entry.get("value", "")

        try:
            if content_type == "table_html":
                return pd.read_html(StringIO(raw_content))[0]

            elif content_type == "html":
                soup = BeautifulSoup(raw_content, "html.parser")
                text_lines = soup.get_text().splitlines()
                text_lines = [line.strip() for line in text_lines if line.strip()]
                return pd.DataFrame({"text": text_lines})

            elif content_type == "pdf":
                pdf_bytes = base64.b64decode(raw_content)
                pdf_file = BytesIO(pdf_bytes)
                all_rows = []
                try:
                    with pdfplumber.open(pdf_file) as pdf:
                        for page_num, page in enumerate(pdf.pages, start=1):
                            tables = page.extract_tables()
                            for table in tables:
                                all_rows.append([f"[Page {page_num}]"])
                                all_rows.extend(table)
                except Exception as e:
                    return pd.DataFrame([[f"Invalid PDF file: {e}"]])
                return pd.DataFrame(all_rows) if all_rows else pd.DataFrame([["No table found in PDF"]])

            elif content_type == "redir_pdf":
                pdf_bytes = base64.b64decode(raw_content)
                pdf_file = BytesIO(pdf_bytes)
                all_rows = []
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_output:
                    ocrmypdf.ocr(pdf_file, temp_output.name)
                try:
                    with pdfplumber.open(temp_output.name) as pdf:
                        for page_num, page in enumerate(pdf.pages, start=1):
                            tables = page.extract_tables()
                            for table in tables:
                                all_rows.append([f"[Page {page_num}]"])
                                all_rows.extend(table)
                except Exception as e:
                    return pd.DataFrame([[f"OCR PDF failed: {e}"]])
                return pd.DataFrame(all_rows) if all_rows else pd.DataFrame([["No table found in OCR PDF"]])

        except Exception as e:
            if hasattr(self, "logger") and self.logger:
                self.logger.error(f"Failed to parse data for content type {content_type}: {e}")
            else:
                print(f"[ERROR] Failed to parse data for content type {content_type}: {e}")
            return pd.DataFrame([[f"Failed to parse data of type: {content_type}"]])
    
    def _write_summary(self, ws, summary, start_row, bank_name="", bank_link=""):
        ws.cell(row=start_row, column=1, value="Bank Name")
        ws.cell(row=start_row, column=2, value=bank_name)

        ws.cell(row=start_row + 1, column=1, value="Bank Link")
        ws.cell(row=start_row + 1, column=2, value=bank_link)

        ws.cell(row=start_row + 2, column=1, value="Summary")
        headers = ["Old Total", "New Total", "New Count", "Removed Count", "Unchanged Count", "Change %"]
        values = [
            summary.get("old_total", 0),
            summary.get("new_total", 0),
            summary.get("new_count", 0),
            summary.get("removed_count", 0),
            summary.get("unchanged_count", 0),
            round((summary.get("new_count", 0) + summary.get("removed_count", 0)) / max(summary.get("old_total", 1), 1) * 100, 2)
        ]

        for col, header in enumerate(headers, start=1):
            ws.cell(row=start_row + 3, column=col, value=header)
        for col, val in enumerate(values, start=1):
            ws.cell(row=start_row + 4, column=col, value=val)

        return start_row + 6
    
    def _write_side_by_side_tables(self, ws, new_df, removed_df, start_row, gap=2):
        from openpyxl.styles import PatternFill
        from openpyxl.formatting.rule import FormulaRule
        from openpyxl.utils.dataframe import dataframe_to_rows

        new_col_start = 1
        removed_col_start = new_df.shape[1] + new_col_start + gap
        comparison_col_start = removed_col_start + removed_df.shape[1] + gap

        SKY_BLUE_FILL = PatternFill(start_color="B3E5FC", end_color="B3E5FC", fill_type="solid")
        GOLD_YELLOW_FILL = PatternFill(start_color="FFE599", end_color="FFE599", fill_type="solid")
        RED_FILL = PatternFill(start_color="FF9999", end_color="FF9999", fill_type="solid")

        max_rows = max(len(new_df), len(removed_df))
        max_cols = max(new_df.shape[1], removed_df.shape[1])

        for r_idx, row in enumerate(dataframe_to_rows(new_df, index=False, header=True)):
            for c_idx, val in enumerate(row):
                cell = ws.cell(row=start_row + r_idx, column=new_col_start + c_idx, value=val)
                cell.fill = SKY_BLUE_FILL

        for r_idx, row in enumerate(dataframe_to_rows(removed_df, index=False, header=True)):
            for c_idx, val in enumerate(row):
                cell = ws.cell(row=start_row + r_idx, column=removed_col_start + c_idx, value=val)
                cell.fill = GOLD_YELLOW_FILL

        for r in range(start_row + 1, start_row + max_rows + 1):
            for c in range(max_cols):
                comp_cell = ws.cell(row=r, column=comparison_col_start + c)
                new_col_letter = ws.cell(row=1, column=new_col_start + c).column_letter
                removed_col_letter = ws.cell(row=1, column=removed_col_start + c).column_letter
                comp_cell.value = f"={new_col_letter}{r}={removed_col_letter}{r}"

        for c in range(max_cols):
            col_letter = ws.cell(row=1, column=comparison_col_start + c).column_letter
            formula = f'{col_letter}{start_row + 1}=FALSE'
            ws.conditional_formatting.add(
                f'{col_letter}{start_row + 1}:{col_letter}{start_row + max_rows}',
                FormulaRule(formula=[formula], fill=RED_FILL)
            )

        return start_row + max_rows + 2
    
    def generate_sorted_excel_report(self, comparison_json, output_path="DepositRate_Comparison_Report.xlsx"):
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.utils.dataframe import dataframe_to_rows
        from openpyxl import load_workbook

        sorted_records = sorted(
            comparison_json.get("records", []),
            key=lambda r: len(r.get("comparison_result", {}).get("new", [])),
            reverse=True
        )

        with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
            for record in sorted_records:
                bank_name = record.get("bank_name")
                bank_code = record.get("bank_code")
                bank_link = record.get("base_url")
                sheet_name = f"{bank_name} ({bank_code})"[:31]

                comparison_result = record.get("comparison_result", {})
                summary = comparison_result.get("summary", {})
                new_entries = comparison_result.get("new", [])
                removed_entries = comparison_result.get("removed", [])

                pd.DataFrame().to_excel(writer, sheet_name=sheet_name, index=False)
                ws = writer.sheets[sheet_name]
                row_cursor = self._write_summary(ws, summary, start_row=1, bank_name=bank_name, bank_link=bank_link)

                max_tables = max(len(new_entries), len(removed_entries))
                for i in range(max_tables):
                    new_df = self._parse_table(new_entries[i]) if i < len(new_entries) else pd.DataFrame()
                    removed_df = self._parse_table(removed_entries[i]) if i < len(removed_entries) else pd.DataFrame()
                    row_cursor = self._write_side_by_side_tables(ws, new_df, removed_df, start_row=row_cursor)

        return output_path
    
    # def _write_summary(self, ws, summary, start_row, bank_name="", bank_link=""):
        
    #     ws.cell(row=start_row, column=1, value="Bank Name")
    #     ws.cell(row=start_row, column=2, value=bank_name)

    #     ws.cell(row=start_row + 1, column=1, value="Bank Link")
    #     ws.cell(row=start_row + 1, column=2, value=bank_link)

    #     ws.cell(row=start_row + 2, column=1, value="Summary")
    #     headers = ["Old Total", "New Total", "New Count", "Removed Count", "Unchanged Count", "Change %"]
    #     values = [
    #         summary.get("old_total", 0),
    #         summary.get("new_total", 0),
    #         summary.get("new_count", 0),
    #         summary.get("removed_count", 0),
    #         summary.get("unchanged_count", 0),
    #         round((summary.get("new_count", 0) + summary.get("removed_count", 0)) / max(summary.get("old_total", 1), 1) * 100, 2)
    #     ]

    #     for col, header in enumerate(headers, start=1):
    #         ws.cell(row=start_row + 3, column=col, value=header)
    #     for col, val in enumerate(values, start=1):
    #         ws.cell(row=start_row + 4, column=col, value=val)

    #     return start_row + 6
    
    # def _parse_table(self, entry):
    #     # print(entry)
    #     content_type = entry.get("type", "str")  # default to html
    #     raw_content = entry.get("value","")
        
    #     try:
    #         if content_type == "table_html":
    #             return pd.read_html(StringIO(raw_content))[0]
            
    #         elif content_type == "html":      
    #             soup = BeautifulSoup(raw_content, "html.parser")
    #             text_lines = soup.get_text().splitlines()
    #             text_lines = [line.strip() for line in text_lines if line.strip()]
    #             return pd.DataFrame({"text": text_lines})
    #         elif content_type == "pdf":
    #             pdf_bytes = base64.b64decode(raw_content)
    #             pdf_file = BytesIO(pdf_bytes)
    #             all_rows = []
    #             with pdfplumber.open(pdf_file) as pdf:
    #                 for page in pdf.pages:
    #                     table = page.extract_table()
    #                     if table:
    #                         all_rows.extend(table)
    #             return pd.DataFrame(all_rows) if all_rows else "[No table found in PDF]"
            
    #         elif content_type == "redir_pdf":
    #             pdf_bytes = base64.b64decode(raw_content)
    #             pdf_file = BytesIO(pdf_bytes)
    #             all_rows = []
    #             with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_output:
    #                 ocrmypdf.ocr(pdf_file, temp_output.name)
    #             with pdfplumber.open(temp_output.name) as pdf:
    #                 for page in pdf.pages:
    #                     tables = page.extract_tables()
    #                     for table in tables:
    #                         all_rows.extend(table)
    #             return pd.DataFrame(all_rows) if all_rows else "[No table found in PDF]"
            

            
    #     except:
    #         self.logger.error(f"Failed to parse data for content type {content_type}")
    #         return pd.DataFrame([[f"Failed to parse data of type: {content_type}"]])
                
    #         # elif content_type == "pdf":
    #         #     pdf_bytes = base64.b64decode(raw_content)
    #         #     pdf_file = BytesIO(pdf_bytes)
    #         #     all_rows = []
    #         #     with pdfplumber.open(pdf_file) as pdf:
    #         #         for page in pdf.pages:
    #         #             table = page.extract_table()
    #         #             if table:
    #         #                 all_rows.extend(table)
    #         #     return pd.DataFrame(all_rows)
            
    #         # elif content_type == "redir_pdf":
    #         #     pdf_bytes = base64.b64decode(raw_content)
    #         #     pdf_file = BytesIO(pdf_bytes)
    #         #     all_rows = []
    #         #     with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_output:
    #         #         ocrmypdf.ocr(pdf_file, temp_output.name)
    #         #     with pdfplumber.open(temp_output.name) as pdf:
    #         #         for page in pdf.pages:
    #         #             tables = page.extract_tables()
    #         #             for table in tables:
    #         #                 all_rows.extend(table)
    #         #     return pd.DataFrame(all_rows)

             
    # def _write_side_by_side_tables(self, ws, new_df, removed_df, start_row, gap=2):
    #     new_col_start = 1
    #     removed_col_start = new_df.shape[1] + new_col_start + gap
    #     comparison_col_start = removed_col_start + removed_df.shape[1] + gap

    #     SKY_BLUE_FILL = PatternFill(start_color="B3E5FC", end_color="B3E5FC", fill_type="solid")
    #     GOLD_YELLOW_FILL = PatternFill(start_color="FFE599", end_color="FFE599", fill_type="solid")
        
    #     max_rows = max(len(new_df), len(removed_df))
    #     max_cols = max(new_df.shape[1], removed_df.shape[1])

    #     # Write new_df
    #     for r_idx, row in enumerate(dataframe_to_rows(new_df, index=False, header=True)):
    #         for c_idx, val in enumerate(row):
    #             cell = ws.cell(row=start_row + r_idx, column=new_col_start + c_idx, value=val)
    #             cell.fill = SKY_BLUE_FILL

    #     # Write removed_df
    #     for r_idx, row in enumerate(dataframe_to_rows(removed_df, index=False, header=True)):
    #         for c_idx, val in enumerate(row):
    #             cell = ws.cell(row=start_row + r_idx, column=removed_col_start + c_idx, value=val)
    #             cell.fill = GOLD_YELLOW_FILL

    #     # Write cell-by-cell comparison formulas
    #     for r in range(start_row + 1, start_row + max_rows + 1):
    #         for c in range(max_cols):
    #             comp_cell = ws.cell(row=r, column=comparison_col_start + c)

    #             # Create formula like =A2=G2
    #             new_col_letter = ws.cell(row=1, column=new_col_start + c).column_letter
    #             removed_col_letter = ws.cell(row=1, column=removed_col_start + c).column_letter
    #             comp_cell.value = f"={new_col_letter}{r}={removed_col_letter}{r}"
                

    #     RED_FILL = PatternFill(start_color="FF9999", end_color="FF9999", fill_type="solid")  # Soft red

    #     #Conditional Formatting 
    #     for c in range(max_cols):
    #         col_letter = ws.cell(row=1, column=comparison_col_start + c).column_letter
    #         formula = f'{col_letter}{start_row + 1}=FALSE'  # Start from first data row
    #         ws.conditional_formatting.add(
    #             f'{col_letter}{start_row + 1}:{col_letter}{start_row + max_rows}',
    #             FormulaRule(formula=[formula], fill=RED_FILL)
    #         )
    #     return start_row + max_rows + 2
    
    # def generate_sorted_excel_report(self, comparison_json, output_path="DepositRate_Comparison_Report.xlsx"):
    #     # Sort records: prioritize those with new entries
    #     sorted_records = sorted(
    #         comparison_json.get("records", []),
    #         key=lambda r: len(r.get("comparison_result", {}).get("new", [])),
    #         reverse=True
    #     )

    #     with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
    #         for record in sorted_records:
    #             bank_name = record.get("bank_name")
    #             bank_code = record.get("bank_code")
    #             bank_link = record.get("base_url")
    #             sheet_name = f"{bank_name} ({bank_code})"[:31]

    #             comparison_result = record.get("comparison_result", {})
    #             summary = comparison_result.get("summary", {})
    #             new_entries = comparison_result.get("new", [])
    #             removed_entries = comparison_result.get("removed", [])

    #             pd.DataFrame().to_excel(writer, sheet_name=sheet_name, index=False)
    #             ws = writer.sheets[sheet_name]
    #             row_cursor = self._write_summary(ws, summary, start_row=1, bank_name=bank_name, bank_link=bank_link)

    #             max_tables = max(len(new_entries), len(removed_entries))
    #             for i in range(max_tables):
    #                 new_df = self._parse_table(new_entries[i]) if i < len(new_entries) else pd.DataFrame()
    #                 removed_df = self._parse_table(removed_entries[i]) if i < len(removed_entries) else pd.DataFrame()
    #                 row_cursor = self._write_side_by_side_tables(ws, new_df, removed_df, start_row=row_cursor)

    #     return output_path
    


    # @staticmethod
    # def __parse_entry(entry):
    #     import fitz  # PyMuPDF

    #     content_type = entry.get("type", "str")
    #     raw_content = entry.get("value", "")
    #     result = {"type": content_type, "content": [], "tables": [], "raw_text": ""}

    #     def clean_cell(cell):
    #         if not cell:
    #             return ""
    #         cleaned = "".join(cell.splitlines())
    #         cleaned = " ".join(cleaned.split())
    #         return cleaned.strip()

    #     try:
    #         if content_type == "pdf":
    #             pdf_bytes = base64.b64decode(raw_content)
    #             pdf_file = BytesIO(pdf_bytes)

    #             with pdfplumber.open(pdf_file) as plumber_pdf:
    #                 fitz_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")

    #                 for page_num, plumber_page in enumerate(plumber_pdf.pages):
    #                     fitz_page = fitz_pdf[page_num]
    #                     elements = []

    #                     # Detect table bounding boxes
    #                     table_bboxes = [table.bbox for table in plumber_page.find_tables()]
    #                     result["tables"].extend([
    #                         pd.DataFrame([
    #                             [clean_cell(cell) for cell in row]
    #                             for row in table.extract()
    #                         ])
    #                         for table in plumber_page.find_tables()
    #                     ])

    #                     # Extract words using fitz
    #                     words = fitz_page.get_text("words")  # [x0, y0, x1, y1, text, ...]

    #                     for w in words:
    #                         x0, y0, x1, y1, text = w[:5]
    #                         top = y0
    #                         in_table = any(
    #                             x0 >= bbox[0] and x1 <= bbox[2] and y0 >= bbox[1] and y1 <= bbox[3]
    #                             for bbox in table_bboxes
    #                         )
    #                         label = "table_text" if in_table else "body_text"
    #                         elements.append({
    #                             "type": "text",
    #                             "text": text,
    #                             "top": top,
    #                             "label": label
    #                         })

    #                     # Extract tables as layout-aware elements
    #                     for table in plumber_page.extract_tables():
    #                         cleaned_table = []
    #                         for row in table:
    #                             cleaned_row = [clean_cell(cell) for cell in row]
    #                             cleaned_table.append(cleaned_row)
    #                         df = pd.DataFrame(cleaned_table)

    #                         # Estimate top position from first word on page
    #                         words_on_page = fitz_page.get_text("words")
    #                         top_pos = float(words_on_page[0][1]) if words_on_page else 0.0

    #                         elements.append({
    #                             "type": "table",
    #                             "table": df,
    #                             "top": top_pos
    #                         })

    #                     # Sort all elements by vertical position
    #                     elements.sort(key=lambda x: float(x.get("top", 0.0)))

    #                     # Build final content stream
    #                     for el in elements:
    #                         if el["type"] == "text":
    #                             result["content"].append({"type": "text", "text": el["text"]})
    #                         elif el["type"] == "table":
    #                             result["content"].append({"type": "table", "table": el["table"]})

    #             result["raw_text"] = "\n".join([
    #                 el["text"] for el in result["content"]
    #                 if el["type"] == "text"
    #             ])

    #         elif content_type == "html":
    #             soup = BeautifulSoup(raw_content, "html.parser")
    #             text = soup.get_text(separator="\n").strip()
    #             result["content"].append({"type": "text", "text": text})
    #             result["raw_text"] = text

    #         elif content_type == "table_html":
    #             df = pd.read_html(StringIO(raw_content))[0]
    #             result["content"].append({"type": "table", "table": df})
    #             result["tables"].append(df)

    #         else:
    #             result["content"].append({"type": "text", "text": f"[Unsupported Datatype: {content_type}]"})

    #     except Exception as e:
    #         result["content"].append({"type": "text", "text": f"[Failed to parse: {e}]"})

    #     return result
    
    # @classmethod
    # def generate_cache_doc_report(cls, comparison_json, output_path="DepositRate_Comparison_Report.docx"):
    #     document = Document()
    #     sorted_records = comparison_json.get("records", [])
    #     metadata = comparison_json.get("metadata", {})

    #     document.add_heading(">>>Metadata Overview<<<", level=1)
    #     table = document.add_table(rows=1, cols=2)
    #     table.style = 'Table Grid'
    #     hdr_cells = table.rows[0].cells
    #     hdr_cells[0].text = 'Key'
    #     hdr_cells[1].text = 'Value'

    #     for key, value in metadata.items():
    #         row_cells = table.add_row().cells
    #         row_cells[0].text = key
    #         row_cells[1].text = value

    #     document.add_paragraph("\n")

    #     for record in sorted_records:
    #         bank_name = record.get("bank_name", "")
    #         bank_code = record.get("bank_code", "")
    #         scraped_data = record.get("scraped_data", [])

    #         document.add_heading(f">> {bank_code} : {bank_name}", level=2)
    #         document.add_paragraph("================================================")

    #         for scrape in scraped_data:
    #             responses = scrape.get("response", [])
    #             action = scrape.get("action", "")
    #             timestamp = scrape.get("timestamp", "")
    #             webpage = scrape.get("webpage", "")
    #             data_present = scrape.get("data_present", "")
    #             response_count = scrape.get("response_count", "")

    #             document.add_heading(
    #                 f"Action: {action} | Timestamp: {timestamp} | Present: {str(data_present)} | Count: {response_count if data_present else 0}",
    #                 level=3
    #             )
    #             document.add_paragraph(f"Website: {webpage}")

    #             for response_entry in responses:
    #                 titles = response_entry.get("title", [])
    #                 titles = [titles] if isinstance(titles, str) else titles

    #                 parsed = cls.__parse_entry(response_entry)
    #                 content_stream = parsed["content"]

    #                 document.add_paragraph("---------------------------")
    #                 for idx, line in enumerate(titles):
    #                     document.add_paragraph(f"TITLE {idx+1}: {line}")

    #                 for item in content_stream:
    #                     if item["type"] == "text":
    #                         document.add_paragraph(item["text"])
    #                     elif item["type"] == "table":
    #                         df = item["table"]
    #                         table = document.add_table(rows=1, cols=len(df.columns))
    #                         table.style = 'Table Grid'
    #                         for i, col_name in enumerate(df.columns):
    #                             table.cell(0, i).text = str(col_name)
    #                         for _, row in df.iterrows():
    #                             row_cells = table.add_row().cells
    #                             for i, val in enumerate(row):
    #                                 row_cells[i].text = str(val)

    #                 document.add_paragraph("")

    #         document.add_page_break()

    #     document.save(output_path)
    #     return output_path
    
    # @classmethod
    # def generate_cache_excel_report(cls, comparison_json, output_path="DepositRate_Comparison_Report.xlsx"):
        # sorted_records = comparison_json.get("records", [])
        # metadata = comparison_json.get("metadata", {})

        # with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        #     meta_df = pd.DataFrame(list(metadata.items()), columns=["Key", "Value"])
        #     meta_df.to_excel(writer, sheet_name="Metadata", index=False)

        #     for record in sorted_records:
        #         bank_name = record.get("bank_name", "UnknownBank")
        #         bank_code = record.get("bank_code", "")
        #         scraped_data = record.get("scraped_data", [])
        #         sheet_name = f"{bank_code}_{bank_name}"[:31]

        #         bank_frames = []

        #         for scrape in scraped_data:
        #             responses = scrape.get("response", [])
        #             action = scrape.get("action", "")
        #             timestamp = scrape.get("timestamp", "")
        #             webpage = scrape.get("webpage", "")
        #             data_present = scrape.get("data_present", "")
        #             response_count = scrape.get("response_count", "")

        #             for response_entry in responses:
        #                 titles = response_entry.get("title", [])
        #                 titles = [titles] if isinstance(titles, str) else titles

        #                 parsed = cls.__parse_entry(response_entry)
        #                 content_stream = parsed["content"]

        #                 header_rows = [
        #                     [f"Action: {action}", f"Timestamp: {timestamp}", f"Present: {data_present}", f"Count: {response_count}"],
        #                     [f"Website: {webpage}"],
        #                     ["---------------------------"],
        #                 ]
        #                 for idx, title in enumerate(titles):
        #                     header_rows.append([f"TITLE {idx+1}: {title}"])

        #                 bank_frames.append(pd.DataFrame(header_rows))

        #                 for item in content_stream:
        #                     if item["type"] == "text":
        #                         bank_frames.append(pd.DataFrame([[item["text"]]]))
        #                     elif item["type"] == "table":
        #                         bank_frames.append(item["table"])

        #                 bank_frames.append(pd.DataFrame([[""]]))  # Spacer

        #         final_df = pd.concat(bank_frames, ignore_index=True)
        #         final_df.to_excel(writer, sheet_name=sheet_name, index=False)

        # return output_path